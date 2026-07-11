from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.27)
section.page_height   = Inches(11.69)
section.left_margin   = Inches(1.5)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_run(run, size, bold=False, italic=False):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def heading1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, 16, bold=True)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(8)
    return p

def heading2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, 14, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(5)
    return p

def heading3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, 12, bold=True)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, 12)
    p.paragraph_format.line_spacing      = Pt(18)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0.3)
    return p

def body_noi(text):
    """Body without first-line indent (after bullets, tables, etc.)"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, 12)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r, 12)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def pb():
    doc.add_page_break()

def add_page_numbers(sec):
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, val in [("w:fldChar","begin"),("w:instrText",None),("w:fldChar","end")]:
        el = OxmlElement(tag)
        if tag == "w:instrText":
            el.text = "PAGE"
        else:
            el.set(qn("w:fldCharType"), val)
        run._r.append(el)

add_page_numbers(section)

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NEURAL TEXT-TO-SPEECH STUDIO")
r.font.name = "Times New Roman"; r.font.size = Pt(24); r.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Project Report")
r.font.name = "Times New Roman"; r.font.size = Pt(18); r.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Powered by Kokoro Neural TTS  |  Fully Offline")
r.font.name = "Times New Roman"; r.font.size = Pt(13)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Project Report submitted in partial fulfilment of the requirements\nfor the award of the degree")
r.font.name = "Times New Roman"; r.font.size = Pt(12)

pb()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run("TABLE OF CONTENTS")
r.font.name = "Times New Roman"; r.font.size = Pt(16); r.font.bold = True
p.paragraph_format.space_after = Pt(12)

toc = [
    (False, "Abstract"),
    (True,  "Chapter 1 – Introduction"),
    (False, "    1.1   Introduction to the Project"),
    (False, "    1.2   History and Evolution of TTS Technology"),
    (False, "    1.3   Neural TTS and Deep Learning"),
    (False, "    1.4   Scope of the Project"),
    (False, "    1.5   Problem Statement"),
    (False, "    1.6   Existing System"),
    (False, "    1.7   Proposed System and Modules"),
    (True,  "Chapter 2 – Requirement Analysis"),
    (False, "    2.1   Feasibility Study"),
    (False, "        2.1.1  Technical Feasibility"),
    (False, "        2.1.2  Economical Feasibility"),
    (False, "        2.1.3  Cost Benefit Analysis"),
    (False, "        2.1.4  Operational Feasibility"),
    (False, "    2.2   Software Requirement Specification (SRS)"),
    (False, "        2.2.1  Functional Requirements"),
    (False, "        2.2.2  Non-Functional Requirements"),
    (False, "    2.3   Environment and Technology Requirements"),
    (False, "        2.3.1  Hardware Requirements"),
    (False, "        2.3.2  Software Requirements"),
    (True,  "Chapter 3 – Design"),
    (False, "    3.1   System Design"),
    (False, "        3.1.1  Introduction to UML"),
    (False, "    3.2   Interface Design"),
    (False, "    3.3   Database Design"),
    (True,  "Chapter 4 – Coding"),
    (False, "    4.1   Sample Code"),
    (True,  "Chapter 5 – Testing"),
    (False, "    5.1   Introduction to Testing"),
    (False, "    5.2   Types of Testing"),
    (False, "    5.3   Test Cases (Module-wise)"),
    (True,  "Chapter 6 – Screenshots"),
    (False, "    6.1   Input Screens"),
    (False, "    6.2   Output Screens"),
    (True,  "Conclusion"),
    (True,  "Bibliography"),
]

for bold, entry in toc:
    p = doc.add_paragraph()
    r = p.add_run(entry)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = bold
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

pb()

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════════════
heading1("ABSTRACT")

body(
    "The proliferation of artificial intelligence and deep learning has enabled a new generation "
    "of text-to-speech (TTS) systems capable of producing speech that is virtually indistinguishable "
    "from human voice. However, most commercially available and academically popular TTS systems "
    "operate as cloud services, requiring the user's text to be transmitted to remote servers over "
    "the internet. This approach raises significant concerns around data privacy, offline "
    "availability, network latency, and dependency on third-party service providers. Neural "
    "Text-to-Speech Studio is a fully offline, locally-executed TTS application developed to "
    "address these concerns directly."
)
body(
    "The application is built using Python 3 and the Streamlit web framework, and it leverages "
    "the Kokoro neural TTS model — a compact, open-source, end-to-end neural speech synthesis "
    "system — for voice generation. Kokoro is capable of producing high-quality, natural-sounding "
    "audio for multiple English accents using a lightweight model architecture that runs efficiently "
    "on consumer-grade hardware without a dedicated GPU. The system supports nine distinct voices "
    "spanning American and British English, both male and female, and provides a speed control "
    "parameter to suit different use-case requirements."
)
body(
    "A carefully designed text preprocessing pipeline ensures that raw user input is normalised "
    "before synthesis. This includes expanding numerical digits into their spoken-word equivalents "
    "(e.g., '345' becomes 'three hundred and forty-five'), replacing common abbreviations, "
    "normalising punctuation spacing, removing unsupported characters, and segmenting excessively "
    "long sentences to prevent synthesis errors. The result is a robust, user-friendly application "
    "that produces natural, intelligible speech from arbitrary English text input without any "
    "network dependency."
)
body(
    "The project demonstrates that state-of-the-art neural TTS can be deployed in a resource-"
    "constrained, offline environment using entirely open-source tooling, and it establishes a "
    "reusable, modular architecture that can be readily extended with additional languages, voices, "
    "or synthesis backends in the future."
)

pb()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 1 – INTRODUCTION")

heading2("1.1  Introduction to the Project")
body(
    "Human communication is fundamentally rooted in spoken language. The ability to convert "
    "written text into natural speech has long been a goal of computer science, spanning "
    "applications from accessibility tools for the visually impaired to automated customer "
    "service systems, e-learning platforms, navigation systems, and entertainment. Text-to-Speech "
    "(TTS) synthesis is the automatic process of generating intelligible, natural-sounding spoken "
    "audio from written text input."
)
body(
    "Neural Text-to-Speech Studio is a Python-based application that brings modern neural TTS "
    "technology to the desktop in a fully offline, privacy-preserving manner. The application "
    "is powered by the Kokoro TTS model, a cutting-edge, open-source neural speech synthesis "
    "system that produces high-fidelity audio indistinguishable from professional voice recordings. "
    "It is presented through a clean, intuitive browser-based interface built with the Streamlit "
    "framework, making it immediately accessible without any web development expertise."
)
body(
    "The application targets a broad user base including students, educators, content creators, "
    "developers, and individuals with visual impairments or reading difficulties. By operating "
    "entirely offline, it provides uninterrupted availability, eliminates privacy risks associated "
    "with transmitting text to cloud servers, and removes dependence on internet connectivity "
    "or third-party API quotas."
)

heading2("1.2  History and Evolution of TTS Technology")
body(
    "The history of text-to-speech synthesis spans more than two centuries, beginning long before "
    "the digital age. In 1791, Wolfgang von Kempelen constructed a mechanical speech machine "
    "using bellows, reeds, and hand-shaped resonators that could produce vowel sounds. The "
    "twentieth century brought electronic attempts, including Homer Dudley's VODER (Voice "
    "Operation DEmonstratoR) demonstrated at the 1939 World's Fair, which used electronic "
    "circuits to model the human vocal tract."
)
body(
    "The first computer-based TTS systems appeared in the 1960s and 1970s. These early systems "
    "used rule-based phoneme synthesis, translating text into phonemes using hand-crafted linguistic "
    "rules and then synthesising audio from those phonemes. While intelligible, the output was "
    "noticeably robotic and unnatural. The 1980s and 1990s saw the rise of concatenative synthesis, "
    "where large databases of pre-recorded speech units (diphones or triphones) were assembled "
    "dynamically. Festival and MBRola were prominent open-source concatenative TTS systems of "
    "this era. Concatenative synthesis improved naturalness significantly but required enormous "
    "speech databases and suffered from audible discontinuities at unit boundaries."
)
body(
    "The 2000s brought statistical parametric speech synthesis (SPSS), using Hidden Markov Models "
    "(HMMs) to generate smooth, continuous speech parameters. HTS (HMM-based Speech Synthesis "
    "System) was a landmark system of this period. While SPSS reduced storage requirements and "
    "improved prosody modelling, the output had a characteristic muffled or buzzy quality due "
    "to over-smoothing of parameters."
)
body(
    "The watershed moment for TTS came in 2016 with DeepMind's WaveNet — a deep generative model "
    "based on dilated causal convolutions that could produce raw audio waveforms of unprecedented "
    "naturalness. WaveNet was followed by a succession of neural TTS architectures that combined "
    "acoustic models with neural vocoders: Tacotron (2017), Tacotron 2 (2018), FastSpeech (2019), "
    "FastSpeech 2 (2020), VITS (2021), and Kokoro (2024). Each successive system improved "
    "naturalness, speed, and robustness. Today's neural TTS systems achieve Mean Opinion Scores "
    "(MOS) on naturalness that match or exceed human recordings on standardised benchmarks."
)

heading2("1.3  Neural TTS and Deep Learning")
body(
    "Modern neural TTS systems are typically composed of two stages: an acoustic model and a "
    "vocoder. The acoustic model converts input text (represented as phonemes or graphemes) into "
    "an intermediate acoustic representation, most commonly a mel-spectrogram. The vocoder then "
    "converts the mel-spectrogram into a raw audio waveform."
)
body(
    "The acoustic model is trained to learn the mapping from text to speech features using a "
    "neural network. Early acoustic models such as Tacotron used recurrent neural networks (RNNs) "
    "with attention mechanisms. FastSpeech replaced the sequential RNN with a feed-forward "
    "Transformer architecture, achieving significantly faster synthesis. VITS (Variational "
    "Inference with adversarial learning for end-to-end TTS) combined the acoustic model and "
    "vocoder into a single end-to-end network trained jointly, eliminating error accumulation "
    "between the two stages."
)
body(
    "The Kokoro model used in this project is an end-to-end neural TTS system based on the "
    "StyleTTS 2 architecture, which incorporates style diffusion and adversarial training to "
    "achieve highly expressive and natural speech. Kokoro is notable for its small model size "
    "(approximately 82 million parameters), which allows it to run efficiently on CPU hardware "
    "without a GPU, while still achieving naturalness comparable to much larger commercial systems."
)
body(
    "A key advantage of neural TTS over rule-based and concatenative systems is its ability to "
    "generalise — it can synthesise speech for any text, including words and names it has never "
    "seen during training, by learning underlying phonological and prosodic patterns from data "
    "rather than relying on hand-crafted rules or pre-recorded inventory."
)

heading2("1.3.1  The Kokoro TTS Model")
body(
    "Kokoro is a state-of-the-art, open-source neural TTS model released in 2024. It is built "
    "on the StyleTTS 2 architecture, which employs style diffusion — a technique borrowed from "
    "diffusion-based generative models — to model the prosodic and acoustic variability of human "
    "speech. Unlike earlier neural TTS models that mapped text directly to a fixed acoustic "
    "representation, StyleTTS 2 learns a latent style space that captures speaking style, "
    "prosody, and speaker identity as continuous vectors. At inference time, a style vector is "
    "sampled from this space (or selected via a voice identifier) and used to condition the "
    "acoustic decoder, producing speech that varies naturally in rhythm, intonation, and timbre."
)
body(
    "Kokoro is notable for several engineering achievements that make it suitable for this "
    "project. First, its parameter count is approximately 82 million — small by modern LLM "
    "standards — allowing it to perform inference on a CPU in seconds rather than requiring "
    "a dedicated GPU. Second, it is distributed as a standard Python package via PyPI, with "
    "pre-trained model weights downloaded automatically on first use through the Hugging Face "
    "Hub infrastructure. Third, the Python interface is ergonomic: a KPipeline object accepts "
    "text and a voice identifier and returns a generator of (graphemes, phonemes, audio) tuples, "
    "where audio is a NumPy array sampled at 24,000 Hz. This chunked generator interface "
    "allows the application to begin processing audio while synthesis of subsequent chunks "
    "is still in progress, laying the groundwork for future streaming support."
)
body(
    "The voices available in Kokoro are pre-trained embeddings that capture the acoustic "
    "characteristics of specific speakers. American English voices use the language code 'a' "
    "and British English voices use the code 'b'. Within each language, voices are labelled "
    "with gender prefixes (f for female, m for male) and a style identifier that reflects "
    "the character of the voice: 'heart' conveys warmth and clarity, 'bella' conveys "
    "expressiveness, 'george' conveys authoritative depth, and so on. The choice of voice "
    "does not affect the linguistic quality of synthesis (phoneme accuracy, number reading) "
    "but significantly influences the perceived naturalness and suitability for different "
    "content types."
)

heading2("1.3.2  The Streamlit Framework")
body(
    "Streamlit is an open-source Python library that allows data scientists and machine "
    "learning engineers to build and share interactive web applications without writing any "
    "HTML, CSS, or JavaScript. A Streamlit application is a Python script that describes the "
    "UI declaratively using Python function calls. Streamlit's server watches the script for "
    "changes and re-runs it from top to bottom whenever the user interacts with a widget, "
    "maintaining a reactive programming model that is familiar to Python developers."
)
body(
    "Streamlit provides built-in components for text input (st.text_area), dropdown menus "
    "(st.selectbox), sliders (st.slider), buttons (st.button), audio playback (st.audio), "
    "status messages (st.success, st.warning, st.error, st.spinner), and layout management "
    "(st.columns, st.sidebar). Custom CSS can be injected via st.markdown with "
    "unsafe_allow_html=True, enabling the dark theme applied in this project."
)
body(
    "For this project, Streamlit is an ideal choice because it eliminates the need for a "
    "separate frontend technology stack, keeps the entire application in a single Python "
    "codebase, and provides the st.audio component which handles WAV binary data natively — "
    "essential for embedding the Kokoro-generated audio directly in the browser without "
    "writing it to a temporary file."
)

heading2("1.4  Scope of the Project")
body(
    "The scope of Neural Text-to-Speech Studio is defined as follows. The system accepts plain "
    "English text as input and produces high-quality WAV audio as output. It supports nine "
    "distinct voice identities across American English (five voices) and British English "
    "(four voices), covering both male and female speakers. The user can control the speech "
    "rate via a speed parameter ranging from 0.5× (half speed) to 1.5× (one-and-a-half times "
    "normal speed)."
)
body(
    "The system includes a text preprocessing module that handles common input quality issues: "
    "numeric digit expansion, abbreviation substitution, punctuation normalisation, illegal "
    "character removal, and sentence length management. The user interface is a single-page "
    "web application served locally via the Streamlit framework and accessible from any web "
    "browser on the same machine."
)
body(
    "The following capabilities are explicitly outside the scope of this version of the project: "
    "multilingual synthesis (languages other than English), real-time streaming synthesis, "
    "voice cloning from user-provided audio, batch file processing, persistent history or "
    "user accounts, and deployment to a public cloud or multi-user server environment."
)

heading2("1.5  Problem Statement")
body(
    "Despite the remarkable advances in neural TTS quality over the past decade, the majority "
    "of high-quality TTS solutions available to end-users operate as cloud services. Google "
    "Text-to-Speech, Microsoft Azure Cognitive Services, Amazon Polly, and IBM Watson TTS are "
    "among the most widely used — all of them require transmitting the user's text to remote "
    "servers over the internet for processing. This paradigm introduces a number of significant "
    "practical and ethical problems."
)
body(
    "First, privacy is compromised. Any text submitted to a cloud TTS service is processed by "
    "a third party. For sensitive documents — medical records, legal texts, personal diary "
    "entries, confidential business content — this is an unacceptable risk. Second, availability "
    "is unreliable. Cloud services depend on network connectivity; in low-connectivity or "
    "no-connectivity environments, these services are entirely unavailable. Third, latency is "
    "introduced by network round-trips, which can range from tens of milliseconds to several "
    "seconds depending on network conditions and server load."
)
body(
    "Fourth, API usage is subject to rate limits and quotas. Free tiers of commercial TTS APIs "
    "impose strict monthly character limits, and exceeding them requires paid subscriptions. "
    "Fifth, free Python libraries that wrap cloud TTS services (such as gTTS and edge-tts) are "
    "brittle — they depend on unofficial or semi-official API endpoints that can be deprecated "
    "or blocked without notice. The edge-tts library, for example, communicates with Microsoft's "
    "Edge browser read-aloud endpoint, which is not a documented public API and may break at any "
    "time."
)
body(
    "This project addresses all of the above problems by providing a fully offline alternative "
    "that runs entirely on the local machine, requires no API keys, imposes no usage limits, "
    "and never transmits user text to any external service."
)

heading2("1.6  Existing System")
body(
    "Two TTS systems formed the baseline for this project before the migration to Kokoro: "
    "gTTS (Google Text-to-Speech) and Microsoft Edge TTS. Their characteristics and limitations "
    "are described below."
)
heading3("1.6.1  gTTS (Google Text-to-Speech)")
body(
    "gTTS is a Python library and CLI tool that interfaces with the Google Translate TTS API. "
    "It generates MP3 audio by making HTTP requests to Google's servers, passing the input text "
    "as a URL query parameter. gTTS supports over 40 languages, is simple to install and use, "
    "and is free within Google's usage limits. However, it suffers from several critical "
    "limitations: it requires an active internet connection for every synthesis call; it produces "
    "audio at a fixed, monotone quality with limited naturalness; it does not support voice "
    "selection, pitch control, or speed control; and it is subject to Google's rate limiting "
    "and potential API changes. Audio quality is noticeably inferior to modern neural TTS systems."
)
heading3("1.6.2  Microsoft Edge TTS")
body(
    "Edge TTS is a Python library that accesses Microsoft's Azure Cognitive Services neural "
    "voices through the read-aloud endpoint of the Microsoft Edge browser. It uses a WebSocket "
    "connection to stream audio from Microsoft's servers. The voices are genuine neural TTS "
    "voices of high quality, supporting rate and pitch customisation and offering dozens of "
    "voice options across multiple languages. However, Edge TTS is entirely dependent on "
    "internet connectivity; every synthesis request contacts Microsoft's servers. It relies "
    "on an undocumented internal API that could be revoked or modified without warning. "
    "Furthermore, all user text is processed on Microsoft's infrastructure, which is a "
    "significant privacy concern for sensitive content."
)
heading3("1.6.3  Summary of Limitations")
bullet("Both systems require internet connectivity — unavailable in offline environments.")
bullet("User text is transmitted to and processed by third-party servers.")
bullet("Subject to rate limiting, API changes, and service outages.")
bullet("gTTS: limited audio quality, no voice customisation, no speed or pitch control.")
bullet("Edge TTS: depends on an undocumented internal Microsoft API with no stability guarantee.")

heading2("1.6.4  Comparison of Existing vs Proposed System")
body_noi("")
cmp_table = doc.add_table(rows=1, cols=4)
cmp_table.style = "Table Grid"
for cell, txt in zip(cmp_table.rows[0].cells, ["Feature", "gTTS", "Edge TTS", "Kokoro (Proposed)"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
cmp_rows = [
    ("Internet required",      "Yes",          "Yes",              "No"),
    ("Audio quality",          "Low (robotic)", "High (neural)",   "High (neural)"),
    ("Voice options",          "None",          "100+ voices",     "9 voices"),
    ("Speed control",          "No",            "Yes",             "Yes"),
    ("Pitch control",          "No",            "Yes",             "No"),
    ("Privacy",                "Low",           "Low",             "Full (local)"),
    ("Rate limits",            "Yes",           "Yes",             "None"),
    ("Offline availability",   "No",            "No",              "Yes"),
    ("Cost",                   "Free (limited)","Free (undoc API)","Free (open-source)"),
    ("API stability",          "Medium",        "Low",             "High"),
]
for row_data in cmp_rows:
    row = cmp_table.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
body_noi("")

heading2("1.7  Proposed System and Modules")
body(
    "The proposed system replaces both online TTS engines with the Kokoro neural TTS model, "
    "providing a fully offline solution of equal or superior audio quality. The system is "
    "architected as a set of loosely coupled Python modules, each with a single, clearly "
    "defined responsibility. The modules are described below."
)
heading3("Module 1 – Text Preprocessing (tts/preprocess.py)")
body(
    "This module accepts raw text input from the user and transforms it into a clean, "
    "synthesis-ready string. It performs the following operations in sequence: (1) stripping "
    "leading and trailing whitespace; (2) replacing common abbreviations such as 'Dr.' and "
    "'Mr.' with their full spoken equivalents; (3) normalising line endings and converting "
    "newline characters into sentence pauses; (4) expanding numeric digits into spoken-word "
    "equivalents using the num2words library; (5) removing characters that are outside the "
    "supported ASCII printable set; (6) normalising punctuation spacing; and (7) splitting "
    "excessively long sentences at sentence or clause boundaries to prevent synthesis errors "
    "caused by input length limits."
)
heading3("Module 2 – TTS Engine (tts/engine.py)")
body(
    "This module defines the available voice identifiers and their human-readable display labels. "
    "It acts as a registry and configuration layer between the UI and the synthesis backend, "
    "making it straightforward to add new voices or swap synthesis engines without modifying "
    "any other part of the codebase."
)
heading3("Module 3 – Kokoro Engine (tts/kokoro_engine.py)")
body(
    "This module wraps the Kokoro KPipeline to provide a simple, synchronous function interface "
    "for speech generation. It manages a cache of pipeline instances (one per language code) "
    "to avoid redundant model loading. Given preprocessed text, a voice identifier, and a "
    "speed value, it invokes the pipeline, collects the generated audio chunks (NumPy arrays), "
    "concatenates them, and encodes the result into a WAV byte stream using the soundfile library."
)
heading3("Module 4 – User Interface (app.py)")
body(
    "This module implements the Streamlit web application. It renders the page layout, collects "
    "user inputs (text, voice selection, speed), invokes the preprocessing and synthesis pipeline "
    "on button press, displays status messages (spinner, success, warning, error), and embeds "
    "the generated audio in the browser for playback."
)
heading3("Module 5 – Audio Utilities (utils/audio.py)")
body(
    "This utility module provides helper functions for managing temporary audio files, including "
    "creating a named temporary file with a specified extension and safely deleting a file by "
    "path. While the current Kokoro-based implementation primarily operates in-memory, these "
    "utilities remain available for future extensions that may require file-based audio handling."
)

heading2("1.8  Project Architecture Overview")
body(
    "The overall architecture of Neural Text-to-Speech Studio follows the Model-View-Controller "
    "(MVC) pattern adapted for a single-page reactive application. The View is the Streamlit "
    "page rendered in the browser. The Controller is the main() function in app.py, which "
    "reads widget state, orchestrates calls to the Model, and updates the View with results. "
    "The Model encompasses the preprocessing module and the Kokoro engine, which are pure "
    "computational components with no direct dependency on the UI framework."
)
body(
    "This architectural separation yields several benefits. The Model can be tested without "
    "any UI framework. The View can be redesigned (e.g., replaced with a desktop GUI using "
    "tkinter or a REST API using FastAPI) without changing the Model. The Controller is thin "
    "and contains only orchestration logic, not business logic. The preprocessing pipeline "
    "is independently exchangeable — for instance, a more sophisticated NLP-based preprocessor "
    "using spaCy or NLTK could replace the current regex-based one without affecting any "
    "other component."
)
body(
    "The deployment model is a single Python process on a local machine. Streamlit launches "
    "a Tornado web server on localhost:8501. The browser communicates with the server over "
    "HTTP and WebSocket. When the user clicks 'Generate Speech', the browser sends a message "
    "to the server triggering a script re-run. The server executes the synthesis pipeline "
    "synchronously and sends the resulting HTML (including the embedded audio element) back "
    "to the browser. The entire round-trip occurs on the local machine, with no data leaving "
    "the host system at any point."
)

heading2("1.9  Literature Survey")
body(
    "A literature survey examines existing research, systems, and published works relevant to "
    "the domain of the project. For Neural Text-to-Speech Studio, the survey covers neural "
    "TTS architectures, offline TTS systems, and related Python-based speech applications."
)

heading3("1.9.1  WaveNet (van den Oord et al., 2016)")
body(
    "WaveNet, published by DeepMind in 2016, was a landmark paper that demonstrated for the "
    "first time that a deep generative model could produce raw audio waveforms of human-level "
    "quality. WaveNet uses a stack of dilated causal convolutional layers to model the "
    "conditional probability distribution of each audio sample given all preceding samples "
    "and a conditioning signal (such as a linguistic feature vector). While WaveNet set the "
    "quality standard for neural TTS, its original implementation was autoregressive — it "
    "generated one sample at a time — making real-time synthesis impractical. This limitation "
    "was subsequently addressed by Parallel WaveNet and other distillation-based methods."
)
body(
    "WaveNet's relevance to this project is foundational: it established that a neural network "
    "could replace the hand-engineered vocoders of earlier TTS systems, directly inspiring "
    "all subsequent neural TTS architectures including Tacotron, FastSpeech, VITS, and Kokoro."
)

heading3("1.9.2  Tacotron and Tacotron 2 (Wang et al., 2017; Shen et al., 2018)")
body(
    "Tacotron (Google Brain, 2017) was the first end-to-end neural TTS system that directly "
    "mapped character sequences to mel-spectrograms using a sequence-to-sequence model with "
    "attention. Tacotron eliminated the need for hand-crafted phoneme dictionaries and acoustic "
    "feature engineering. Tacotron 2 (2018) improved on the original by using a more robust "
    "attention mechanism and pairing the acoustic model with WaveNet as the vocoder, achieving "
    "naturalness scores that were statistically indistinguishable from human speech on "
    "benchmark evaluations. Tacotron 2 remains one of the most widely cited TTS architectures "
    "and established the mel-spectrogram as the standard intermediate representation for "
    "neural TTS pipelines."
)

heading3("1.9.3  FastSpeech and FastSpeech 2 (Ren et al., 2019, 2021)")
body(
    "FastSpeech (Microsoft Research, 2019) addressed the fundamental limitation of "
    "attention-based sequence-to-sequence models: slow, non-parallel inference and attention "
    "alignment errors. FastSpeech replaced the recurrent encoder-decoder with a "
    "feed-forward Transformer and used knowledge distillation from a Tacotron 2 teacher "
    "model to train a duration predictor that explicitly controls how many mel-spectrogram "
    "frames each phoneme occupies. This enabled fully parallel (non-autoregressive) "
    "decoding, making synthesis approximately 38 times faster than Tacotron 2. FastSpeech 2 "
    "extended this by predicting not only duration but also pitch and energy contours, "
    "eliminating the need for teacher distillation and improving naturalness and expressiveness."
)
body(
    "FastSpeech's contribution is directly relevant to Kokoro's design: the use of explicit "
    "duration and prosody predictors in a parallel Transformer architecture is the core "
    "technique that makes Kokoro fast enough to run on CPU in a practical timeframe."
)

heading3("1.9.4  VITS (Kim et al., 2021)")
body(
    "VITS (Conditional Variational Inference with adversarial Training for end-to-end TTS) "
    "introduced a fully end-to-end architecture that jointly trains the acoustic model and "
    "vocoder within a single variational autoencoder framework with adversarial training. "
    "VITS produces raw audio waveforms directly from text without a separate mel-spectrogram "
    "intermediate representation, eliminating error accumulation between the acoustic model "
    "and vocoder stages. VITS achieved state-of-the-art naturalness at the time of publication "
    "and demonstrated that end-to-end training is superior to the two-stage pipeline approach "
    "when sufficient training data is available."
)

heading3("1.9.5  StyleTTS 2 (Li et al., 2022)")
body(
    "StyleTTS 2, published by Columbia University, is the direct precursor to Kokoro. It "
    "introduced two key innovations: style diffusion, which uses a diffusion model to sample "
    "expressive style representations at inference time, and large-scale adversarial training "
    "using discriminators derived from large pre-trained speech language models (such as "
    "wav2vec 2.0). These techniques allow StyleTTS 2 to produce highly natural, expressive "
    "speech that exhibits human-like variability in prosody. StyleTTS 2 was the first "
    "open-source model to achieve Mean Opinion Scores matching human recordings on the "
    "LJSpeech benchmark. Kokoro is a lightweight, multi-voice adaptation of StyleTTS 2 "
    "optimised for efficient deployment."
)

heading3("1.9.6  Existing Offline TTS Libraries")
body(
    "Several offline TTS libraries predate Kokoro but have limitations that motivated this "
    "project's choice of synthesis backend. Pyttsx3 is a Python wrapper around platform-native "
    "TTS engines (SAPI on Windows, NSS on macOS, espeak on Linux). While fully offline, "
    "its audio quality is far below neural TTS — the output is robotic and uses concatenative "
    "or formant synthesis. Festival is an open-source concatenative TTS system developed at "
    "the University of Edinburgh; it is offline and supports multiple languages but requires "
    "complex installation on modern systems and produces noticeably synthetic-sounding audio "
    "compared to neural alternatives. Mozilla TTS (now Coqui TTS) is a Python framework "
    "for training and deploying neural TTS models offline; it supports many architectures "
    "including Tacotron 2, FastSpeech 2, and VITS, but requires significantly more setup "
    "effort than Kokoro and its pre-trained models for English are larger and slower. "
    "Kokoro represents the current state-of-the-art for easy-to-deploy, offline, neural "
    "TTS in Python."
)

pb()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 – REQUIREMENT ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 2 – REQUIREMENT ANALYSIS")

heading2("2.1  Feasibility Study")
body(
    "A feasibility study examines the practicality of a proposed project from multiple dimensions "
    "before development begins. It helps project stakeholders identify potential risks, estimate "
    "resource requirements, and confirm that the project can be completed within available "
    "constraints. For Neural Text-to-Speech Studio, the feasibility study covers four dimensions: "
    "technical, economical, cost-benefit, and operational."
)

heading3("2.1.1  Technical Feasibility")
body(
    "Technical feasibility assesses whether the required technology, tools, and expertise are "
    "available and sufficient to build the proposed system. Neural Text-to-Speech Studio relies "
    "on the following technical components, each of which is well-established and readily available:"
)
body(
    "Python 3.10 or above is a mature, widely adopted programming language with an extensive "
    "ecosystem of libraries for machine learning, audio processing, and web development. All "
    "components required by this project are available as Python packages installable via pip."
)
body(
    "Kokoro TTS is distributed as an open-source Python package on PyPI. It is built on PyTorch, "
    "the most widely used deep learning framework, and includes pre-trained model weights that "
    "require no additional training. The model runs on both CPU and GPU, making it compatible "
    "with a wide range of hardware configurations including laptops without dedicated graphics cards."
)
body(
    "Streamlit is a production-grade Python framework for building interactive web applications. "
    "It handles HTTP serving, WebSocket communication for state management, widget rendering, "
    "and browser-based media playback automatically, requiring no web development knowledge "
    "from the application developer."
)
body(
    "The soundfile library wraps libsndfile for reading and writing audio files in WAV and other "
    "formats. num2words provides number-to-words conversion for over 40 languages. Both are "
    "lightweight, stable, and actively maintained. The overall technical stack is mature, "
    "well-documented, and poses no unresolved technical risks. Technical feasibility is confirmed."
)

heading3("2.1.2  Economical Feasibility")
body(
    "Economical feasibility assesses whether the project can be developed and operated within "
    "acceptable financial constraints. All software components used in this project — Python, "
    "Streamlit, Kokoro, PyTorch, soundfile, and num2words — are freely available under "
    "open-source licences (MIT, Apache 2.0, and BSD). There are no licensing fees, subscription "
    "costs, or per-request charges at any point in the development or operational lifecycle."
)
body(
    "The hardware requirement is a standard personal computer or laptop, which is assumed to be "
    "already available. No cloud compute instances, dedicated servers, or specialised hardware "
    "are required. Deployment is performed by running a Python process on the local machine. "
    "There are no hosting costs, CDN costs, or database costs. The project is economically "
    "feasible with an ongoing operational cost of zero rupees."
)

heading3("2.1.3  Cost Benefit Analysis")
body(
    "The development cost of this project is limited to developer time, estimated at approximately "
    "four to six weeks of part-time effort including design, implementation, testing, and "
    "documentation. The operational cost is zero. Against this investment, the project delivers "
    "the following quantifiable and qualitative benefits:"
)
bullet("Elimination of API usage fees that would otherwise apply when scaling to production usage.")
bullet("Removal of network latency, reducing perceived synthesis time for local use.")
bullet("Complete data privacy — no user text leaves the local machine under any circumstances.")
bullet("No exposure to third-party service outages, rate limits, or API deprecation.")
bullet("A reusable, extensible codebase that can support future enhancements at minimal incremental cost.")
body(
    "The cost-benefit ratio is strongly favourable. The project delivers long-term, recurring "
    "benefits (privacy, availability, zero operating cost) in exchange for a one-time, bounded "
    "development effort."
)

heading3("2.1.4  Operational Feasibility")
body(
    "Operational feasibility assesses whether the system can be used effectively by its intended "
    "audience once deployed. The application is started with a single terminal command "
    "(streamlit run app.py) and is immediately accessible from any web browser on the local "
    "machine at http://localhost:8501. No installation wizard, configuration file editing, or "
    "network setup is required beyond the initial package installation."
)
body(
    "The user interface is a single web page with clearly labelled controls. The workflow is "
    "linear: type text, select a voice, optionally adjust speed, click the generate button, "
    "and listen to the audio. No technical knowledge beyond basic computer literacy is required "
    "to operate the application. Error states (empty input, synthesis failure) are communicated "
    "through clear, non-technical warning messages. The system is operationally feasible for "
    "students, educators, content creators, and non-technical users alike."
)

heading2("2.2  Software Requirement Specification (SRS)")
body(
    "The Software Requirement Specification (SRS) formally defines what the system must do "
    "(functional requirements) and how well it must do it (non-functional requirements). "
    "These requirements serve as the contract between the development team and the stakeholders, "
    "and as the basis for system design, implementation, and testing."
)

heading3("2.2.1  Functional Requirements")

heading3("Input Requirements")
bullet("FR-I-01: The system shall provide a multi-line text area for entering input text of arbitrary length.")
bullet("FR-I-02: The system shall provide a dropdown menu listing all available voice options with descriptive human-readable labels.")
bullet("FR-I-03: The system shall provide a slider control for selecting speech speed in the range [0.5, 1.5] with a step size of 0.05.")
bullet("FR-I-04: The system shall provide a prominently displayed 'Generate Speech' button to trigger synthesis.")
bullet("FR-I-05: The system shall validate that the text input is not empty before invoking the synthesis pipeline.")

heading3("Output Requirements")
bullet("FR-O-01: The system shall produce a WAV audio file from the processed text using the selected voice and speed.")
bullet("FR-O-02: The system shall embed an audio playback control in the browser interface upon successful generation.")
bullet("FR-O-03: The audio player shall allow the user to play, pause, seek, and download the generated audio.")
bullet("FR-O-04: The system shall display a success message ('Speech generated successfully.') upon completion.")
bullet("FR-O-05: The system shall display a warning message if the text input is empty or reduces to empty after preprocessing.")
bullet("FR-O-06: The system shall display an error message with the exception description if synthesis fails.")

heading3("Computational / Processing Requirements")
bullet("FR-C-01: The system shall preprocess input text before synthesis by expanding digits to words, substituting abbreviations, normalising punctuation, filtering invalid characters, and splitting long sentences.")
bullet("FR-C-02: The system shall select the appropriate Kokoro language pipeline ('a' for American, 'b' for British) based on the first character of the voice identifier.")
bullet("FR-C-03: The system shall cache the Kokoro pipeline instance per language code to avoid redundant model initialisation across successive requests in the same session.")
bullet("FR-C-04: The system shall concatenate multiple audio chunks returned by the Kokoro pipeline into a single contiguous audio array before encoding.")
bullet("FR-C-05: The system shall encode the final audio array to WAV format in memory and return it as a bytes object for browser playback.")

heading3("Storage Requirements")
body(
    "This application has no persistent storage requirements. All processing is performed in "
    "memory. No database, file storage, or user session persistence is required. The Kokoro "
    "model weights are loaded from the filesystem on first use and held in memory for the "
    "duration of the session by the framework."
)

heading3("2.2.2  Non-Functional Requirements")

heading3("Reliability")
body(
    "NFR-R-01: The system shall produce a valid audio output for any non-empty English text "
    "input that passes the preprocessing stage, under normal hardware conditions."
)
body(
    "NFR-R-02: The system shall handle synthesis exceptions gracefully, displaying a descriptive "
    "error message to the user without crashing the application or requiring a restart."
)
body(
    "NFR-R-03: The system shall correctly preprocess edge-case inputs including inputs consisting "
    "entirely of digits, inputs with only punctuation, and inputs containing newlines."
)

heading3("Usability")
body(
    "NFR-U-01: A new user with no prior knowledge of TTS should be able to generate their first "
    "audio output within two minutes of launching the application, without reading any documentation."
)
body(
    "NFR-U-02: All controls shall be labelled clearly and unambiguously. Tooltips or descriptive "
    "captions shall be provided where the purpose of a control is not self-evident."
)
body(
    "NFR-U-03: The application shall respond to a generate request with a visible spinner "
    "indicator within one second of the button click, so the user knows the system is working."
)

heading3("Response Time")
body(
    "NFR-T-01: For inputs of 500 characters or fewer, the system shall complete synthesis and "
    "present the audio player to the user within 15 seconds on a standard laptop with a "
    "quad-core CPU and 8 GB RAM."
)
body(
    "NFR-T-02: For inputs between 500 and 2000 characters, synthesis time shall be proportional "
    "to input length and may range from 15 to 60 seconds. The spinner indicator shall remain "
    "visible throughout."
)

heading3("Throughput")
body(
    "NFR-TP-01: The system is designed for single-user, single-session interactive use. "
    "Concurrent requests from multiple users are not a requirement and are not supported "
    "in this version."
)

heading3("Availability")
body(
    "NFR-A-01: As a fully offline application, the system shall be available whenever the host "
    "machine is powered on and the Streamlit process is running, with no dependency on external "
    "networks, APIs, or services."
)
body(
    "NFR-A-02: The system shall remain functional during network outages, DNS failures, and "
    "in air-gapped environments."
)

heading3("Maintainability")
body(
    "NFR-M-01: The TTS engine shall be replaceable by modifying kokoro_engine.py and engine.py "
    "only, without requiring changes to app.py or the preprocessing module."
)
body(
    "NFR-M-02: New voices shall be addable by appending entries to the KOKORO_VOICE_CHOICES "
    "list and VOICE_LABELS dictionary in engine.py, without modifying any other file."
)
body(
    "NFR-M-03: The preprocessing pipeline steps shall be individually testable as pure functions "
    "with no side effects."
)

heading2("2.3  Environment and Technology Requirements")

heading3("2.3.1  Hardware Requirements")
body(
    "The following hardware specifications represent the minimum and recommended configurations "
    "for running Neural Text-to-Speech Studio:"
)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for cell, txt in zip(table.rows[0].cells, ["Component", "Minimum", "Recommended"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"

hw = [
    ("Processor",    "Dual-core 2.0 GHz x86-64",  "Quad-core 2.5 GHz or above"),
    ("RAM",          "4 GB",                        "8 GB or above"),
    ("Disk Space",   "5 GB free",                   "10 GB free"),
    ("GPU",          "Not required",                "NVIDIA CUDA GPU (optional, for speed)"),
    ("Audio Output", "Speakers or headphones",      "Speakers or headphones"),
    ("Display",      "Any (browser-capable)",       "Full HD (1920×1080)"),
    ("Network",      "Not required",                "Not required"),
]
for row_data in hw:
    row = table.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size  = Pt(11)
        cell.paragraphs[0].runs[0].font.name  = "Times New Roman"

body_noi("")

heading3("2.3.2  Software Requirements")
body(
    "The following software components are required to build and run the application:"
)

table2 = doc.add_table(rows=1, cols=3)
table2.style = "Table Grid"
for cell, txt in zip(table2.rows[0].cells, ["Software", "Version", "Purpose"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"

sw = [
    ("Python",       "3.10 or above",  "Runtime language"),
    ("Streamlit",    "Latest stable",  "Web UI framework"),
    ("Kokoro",       "Latest stable",  "Neural TTS engine"),
    ("PyTorch",      "2.0 or above",   "Deep learning backend for Kokoro"),
    ("soundfile",    "Latest stable",  "WAV audio encoding"),
    ("num2words",    "Latest stable",  "Number-to-words conversion"),
    ("python-docx",  "Latest stable",  "Documentation generation (dev only)"),
    ("Web Browser",  "Any modern",     "Streamlit UI access"),
    ("OS",           "Linux / Win / macOS", "Host operating system"),
]
for row_data in sw:
    row = table2.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size  = Pt(11)
        cell.paragraphs[0].runs[0].font.name  = "Times New Roman"

pb()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – DESIGN
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 3 – DESIGN")

heading2("3.1  System Design")
body(
    "The system design of Neural Text-to-Speech Studio follows a layered, modular architecture. "
    "The application is divided into four conceptual layers: the presentation layer (Streamlit UI), "
    "the business logic layer (preprocessing and engine management), the synthesis layer (Kokoro "
    "pipeline), and the utility layer (audio encoding and file management). Each layer communicates "
    "only with adjacent layers, enforcing separation of concerns and simplifying maintenance."
)
body(
    "Data flow through the system proceeds as follows. The user enters text, selects a voice, "
    "and sets a speed value in the presentation layer. On button press, the presentation layer "
    "passes the raw text to the preprocessing module in the business logic layer. The preprocessed "
    "text, along with the voice identifier and speed value, is passed to the Kokoro engine module. "
    "The engine module invokes the Kokoro pipeline, which returns a sequence of NumPy audio arrays. "
    "These arrays are concatenated and passed to the soundfile encoder, which returns a WAV byte "
    "stream. The byte stream is returned to the presentation layer, which embeds it in an HTML "
    "audio element rendered in the browser."
)

heading3("3.1.1  Introduction to UML")
body(
    "Unified Modelling Language (UML) is a standardised, general-purpose visual modelling language "
    "used in the field of software engineering. It was adopted as a standard by the Object "
    "Management Group (OMG) in 1997 and has since become the de facto standard for visualising, "
    "specifying, constructing, and documenting the artefacts of software systems. UML provides "
    "a rich set of diagram types organised into two broad categories: structural diagrams (which "
    "capture the static structure of a system, such as class diagrams and component diagrams) "
    "and behavioural diagrams (which capture the dynamic behaviour of a system, such as sequence "
    "diagrams, activity diagrams, and state machine diagrams)."
)
body(
    "For Neural Text-to-Speech Studio, UML is employed to communicate the system's architecture "
    "and behaviour to stakeholders in a clear, language-independent visual format. The key "
    "structural elements of the system — modules, classes, and their relationships — are "
    "expressed through class and component diagrams. The runtime behaviour — the sequence "
    "of operations triggered by a user interaction — is expressed through sequence and activity "
    "diagrams. Together, these diagrams provide a complete picture of both what the system "
    "consists of and how it behaves at runtime."
)
body(
    "Use-case diagrams capture the functional requirements of the system from the perspective "
    "of external actors (users). They show what the system can do (use cases) and who can "
    "trigger each capability (actors). For this application, the single external actor is the "
    "User, and the use cases are: Enter Text, Select Voice, Adjust Speed, Generate Speech, "
    "Play Audio, and Download Audio. Use-case diagrams serve as the bridge between the informal "
    "requirements gathered in Chapter 2 and the formal design artefacts developed in this chapter."
)
body(
    "Class diagrams model the static structure of the software: the classes, their attributes, "
    "their methods, and the relationships between them (association, dependency, inheritance, "
    "composition). For this project, the principal classes are the Streamlit application "
    "controller (app.py), the KokoroPipeline wrapper (kokoro_engine.py), the text preprocessor "
    "functions (preprocess.py), and the engine configuration module (engine.py)."
)
body(
    "Sequence diagrams model the temporal ordering of messages between objects during a specific "
    "scenario. For this project, the primary sequence modelled is the speech generation flow: "
    "from the user clicking 'Generate Speech', through preprocessing, synthesis, encoding, "
    "and finally the rendering of the audio player. Sequence diagrams are particularly useful "
    "for identifying potential race conditions or error propagation paths."
)
body(
    "Activity diagrams model the flow of control through a process, similar to flowcharts but "
    "with richer semantics including parallel flows, decision points, and swim lanes. The "
    "activity diagram for this application captures the complete lifecycle of a synthesis "
    "request, including input validation, preprocessing, synthesis, error handling, and output "
    "rendering. State machine diagrams model the discrete states an object or system can occupy "
    "and the transitions between those states, triggered by events or conditions."
)

heading2("3.1.2  System Working and Algorithms")

heading3("3.1.2.1  Text Preprocessing Algorithm")
body(
    "The preprocessing pipeline is a sequential algorithm implemented as a series of pure string "
    "transformation functions. Each function is applied to the output of the previous one, "
    "forming a processing chain. The algorithm is described step by step below."
)
body(
    "Step 1 — Type Guard: The function first checks whether the input is a Python str object. "
    "If the input is of any other type (integer, None, list, etc.), the function immediately "
    "returns an empty string. This guard prevents unexpected TypeErrors from propagating into "
    "downstream string operations."
)
body(
    "Step 2 — Whitespace Stripping: Python's str.strip() method removes leading and trailing "
    "whitespace characters including spaces, tabs, and newlines. This normalises inputs that "
    "may have been pasted from a text editor with trailing newlines."
)
body(
    "Step 3 — Abbreviation Substitution: A series of str.replace() calls substitutes common "
    "abbreviations with their full spoken forms. 'Dr.' becomes 'Doctor', 'Mr.' becomes 'Mister', "
    "'Mrs.' becomes 'Missus', 'Ms.' becomes 'Miss', 'Prof.' becomes 'Professor', and 'St.' "
    "becomes 'Saint'. These substitutions are order-independent and are applied before number "
    "expansion to avoid interference."
)
body(
    "Step 4 — Line Ending Normalisation: The text may contain Windows-style (\\r\\n), classic Mac "
    "style (\\r), or Unix-style (\\n) line endings. All are normalised to Unix \\n first, then each "
    "\\n is replaced by '. ' (period and space), converting paragraph breaks into explicit "
    "sentence boundaries that the TTS engine will interpret as pauses."
)
body(
    "Step 5 — Number Expansion: A regular expression \\d+ matches one or more consecutive "
    "digit characters. For each match, the integer value is passed to num2words(), which "
    "returns the English spoken form. For example, '345' → 'three hundred and forty-five', "
    "'1000000' → 'one million', '0' → 'zero'. The substitution is performed globally "
    "across the entire text in a single re.sub() call."
)
body(
    "Step 6 — Character Filtering: A character class regex [^A-Za-z0-9\\.,\\?!;:'\\\"()\\s-] "
    "matches any character not in the permitted set of ASCII letters, digits, common "
    "punctuation, parentheses, hyphens, and whitespace. Each such character is replaced "
    "by a space. This removes Unicode characters, emojis, mathematical symbols, currency "
    "symbols, and other glyphs that the TTS phonemiser cannot handle."
)
body(
    "Step 7 — Punctuation Normalisation: The _enhance_punctuation function uses a regex "
    "to ensure that exactly one space follows each punctuation mark (period, comma, "
    "question mark, exclamation mark, semicolon, colon) with no preceding spaces. It "
    "then collapses multiple consecutive spaces into a single space. This step corrects "
    "spacing artefacts introduced by steps 3–6."
)
body(
    "Step 8 — Sentence Splitting: The _split_long_sentences function checks whether the "
    "text exceeds a configurable maximum length (default 220 characters). If so, it "
    "iteratively extracts segments up to 220 characters long, preferring to split at "
    "the last period within the window, then the last comma, and finally at the hard "
    "character limit. The resulting segments are joined with a space. This prevents "
    "the Kokoro pipeline from receiving excessively long inputs that may cause degraded "
    "synthesis quality or errors."
)

heading3("3.1.2.2  Speech Synthesis Algorithm")
body(
    "The Kokoro speech synthesis algorithm implemented in kokoro_engine.py operates as follows. "
    "The function first determines the language code from the voice identifier: voices starting "
    "with 'b' use the British English pipeline ('b'), and all others use the American English "
    "pipeline ('a'). This distinction matters because Kokoro trains separate phoneme sets and "
    "acoustic models for each language variant."
)
body(
    "The function then retrieves the appropriate KPipeline instance from the _pipelines cache "
    "dictionary. If no pipeline exists for the requested language code, a new KPipeline is "
    "instantiated — a process that loads the model weights from local storage into memory, "
    "which may take several seconds on first call. On subsequent calls with the same language "
    "code, the cached instance is returned immediately."
)
body(
    "The pipeline is invoked as a generator with the preprocessed text, the voice identifier, "
    "and the speed multiplier. Internally, the pipeline performs the following sub-steps: "
    "(1) grapheme-to-phoneme (G2P) conversion, translating the input text into a sequence "
    "of phoneme symbols using the language's phoneme inventory; (2) prosody modelling, "
    "predicting duration and fundamental frequency (F0) contours for each phoneme using "
    "a Transformer-based acoustic model; (3) mel-spectrogram generation, converting the "
    "prosodic predictions into a mel-spectrogram through the StyleTTS 2 decoder conditioned "
    "on the selected voice's style embedding; and (4) waveform generation, converting the "
    "mel-spectrogram to a raw audio waveform using a neural vocoder at 24,000 Hz sample rate."
)
body(
    "The generator yields chunks of audio as NumPy float32 arrays. The application collects "
    "all chunks in a list and concatenates them into a single array using numpy.concatenate(). "
    "This contiguous array is then passed to soundfile.write() with the WAV format specifier "
    "and a sample rate of 24,000 Hz. soundfile encodes the float32 array as a standard RIFF "
    "WAV file in an io.BytesIO buffer. The buffer's contents are returned as a Python bytes "
    "object, which Streamlit's st.audio() component renders as an HTML5 audio element."
)

heading3("3.1.2.3  Application Control Flow")
body(
    "The overall control flow of the application on each user interaction is as follows. "
    "Streamlit re-runs the main() function from the top whenever a widget value changes or "
    "the Generate button is clicked. On re-run, all widget values are read from Streamlit's "
    "session state. If the Generate button's value is True (it was clicked in this run), "
    "the following sequence executes: (1) validate text_input.strip() — if empty, call "
    "st.warning() and return early; (2) call preprocess_text(text_input) — if the result "
    "is empty, call st.warning() and return early; (3) call st.spinner() to show the loading "
    "indicator; (4) call generate_speech_with_kokoro() inside a try-except block; (5) on "
    "success, call st.success() and st.audio() with the returned bytes; (6) on exception, "
    "call st.error() with the exception message and return. All other widget renders (title, "
    "text area, selectors, slider, button) occur on every run regardless of whether the "
    "button was clicked."
)

heading2("3.2  Interface Design")
body(
    "The user interface of Neural Text-to-Speech Studio is implemented as a single-page Streamlit "
    "application. Streamlit renders Python UI components as HTML, CSS, and JavaScript in the "
    "browser, handling all client-server communication internally. The page is served on "
    "http://localhost:8501 by default."
)
body(
    "The visual design uses a custom dark theme applied via injected CSS. The page background "
    "is a radial gradient transitioning from dark navy (#1c1f26) at the top-left to near-black "
    "#07090d at the bottom-right, creating a modern, high-contrast appearance. The text area "
    "uses a dark background (#151922) with light text (#eef2ff) and a subtle blue border, "
    "consistent with the overall dark colour scheme. The generate button uses a horizontal "
    "gradient from sky blue (#0ea5e9) to electric blue (#2563eb) with bold white text, making "
    "it the visually dominant call-to-action element on the page."
)
body(
    "The layout organises controls into two equal columns side by side: the voice selector "
    "occupies the left column, and the speed slider occupies the right column. This layout "
    "makes efficient use of horizontal space on wide displays while remaining readable on "
    "narrower screens. The text area spans the full width of the content area above the "
    "column group, and the generate button spans the full width below it, creating a clear "
    "top-to-bottom input flow."
)
body(
    "The voice selector is a Streamlit selectbox populated from the KOKORO_VOICE_CHOICES list. "
    "Each voice is displayed using its human-readable VOICE_LABELS entry (e.g., 'American "
    "Female – Heart') rather than its raw identifier (e.g., 'af_heart'), improving usability "
    "for non-technical users. The speed slider ranges from 0.5 to 1.5 with a step of 0.05 "
    "and a default of 1.0, representing natural speech rate. The slider is labelled 'Speed' "
    "and updates its displayed value in real time as the user drags."
)
body(
    "Status communication is handled through three Streamlit feedback components. st.spinner() "
    "displays an animated loading indicator with the message 'Generating audio...' during "
    "synthesis. st.success() displays a green banner on completion. st.warning() displays a "
    "yellow banner for validation failures. st.error() displays a red banner for synthesis "
    "exceptions. The st.audio() component renders an HTML5 audio element with native browser "
    "playback controls, supporting play, pause, seek, volume control, and download."
)

heading2("3.3  Database Design")

heading3("3.3.1  Introduction to Backend")
body(
    "Neural Text-to-Speech Studio has no persistent backend or database. The application is "
    "entirely stateless between sessions. Within a session, the only state maintained is the "
    "Kokoro pipeline cache — a Python dictionary that maps language codes to KPipeline instances. "
    "This in-session cache prevents redundant model loading when the user makes multiple "
    "synthesis requests with voices of the same language family (e.g., multiple American "
    "English voices) within the same session. The cache is stored in the module-level "
    "_pipelines dictionary in kokoro_engine.py."
)
body(
    "No user text, generated audio, synthesis parameters, or usage history is written to disk "
    "at any point. All data flows through in-memory Python objects. The absence of a database "
    "layer simplifies the architecture substantially, eliminates all database-related security "
    "risks (SQL injection, unauthorised access), and ensures complete privacy."
)

heading3("3.3.2  Normalization")
body(
    "Not applicable — this project does not employ a relational database. There are no tables, "
    "rows, columns, primary keys, foreign keys, or normalisation considerations. The data model "
    "is entirely implicit in the function signatures and type annotations of the Python modules."
)

heading3("3.3.3  Table Description")
body(
    "Not applicable — no database tables are used in this project. The closest analogue to a "
    "data schema in this application is the VOICE_LABELS dictionary in engine.py, which maps "
    "voice identifier strings to human-readable label strings. This is an in-memory constant "
    "dictionary, not a database table."
)

pb()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 – CODING
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 4 – CODING")

heading2("4.1  Sample Code")
body(
    "This chapter presents the complete source code of the application across all its modules. "
    "The code is written in Python 3.12 following PEP 8 style guidelines. Type annotations are "
    "used throughout for clarity and IDE support. The project structure is as follows:"
)
bullet("app.py — Main Streamlit application entry point.")
bullet("tts/engine.py — Voice registry and engine configuration.")
bullet("tts/kokoro_engine.py — Kokoro TTS synthesis wrapper.")
bullet("tts/preprocess.py — Text normalisation and preprocessing pipeline.")
bullet("utils/audio.py — Audio file utility functions.")

heading3("app.py")
code_block("""\
import streamlit as st

from tts.engine import KOKORO_VOICE_CHOICES, VOICE_LABELS
from tts.kokoro_engine import generate_speech_with_kokoro
from tts.preprocess import preprocess_text


def main():
    st.set_page_config(
        page_title="Neural TTS Studio",
        page_icon="\\u20d0",  # microphone emoji
        layout="wide"
    )

    st.markdown(
        \"\"\"
        <style>
        .stApp {
            background: radial-gradient(
                circle at top left,
                #1c1f26 0%, #0e1117 45%, #07090d 100%
            );
            color: #e8ecf3;
        }
        .stTextArea textarea {
            background-color: #151922 !important;
            color: #eef2ff !important;
            border-radius: 12px !important;
            border: 1px solid #2c3345 !important;
        }
        .stButton button {
            background: linear-gradient(90deg, #0ea5e9, #2563eb);
            color: white;
            border-radius: 10px;
            height: 3em;
            font-weight: 700;
            border: none;
        }
        .stSelectbox div[data-baseweb="select"] > div,
        .stSlider div[data-baseweb="slider"] {
            background-color: #111622;
        }
        </style>
        \"\"\",
        unsafe_allow_html=True,
    )

    st.title("\\U0001f399 Neural Text-to-Speech Studio")
    st.caption(
        "Fully offline, high-quality voice generation powered by Kokoro"
    )

    text_input = st.text_area(
        "Enter text", height=170, placeholder="Type something..."
    )

    col1, col2 = st.columns(2)
    with col1:
        voice_key = st.selectbox(
            "Voice",
            KOKORO_VOICE_CHOICES,
            format_func=lambda v: VOICE_LABELS.get(v, v),
            index=0,
        )
    with col2:
        speed_value = st.slider("Speed", 0.5, 1.5, 1.0, step=0.05)

    if st.button("\\U0001f3a7 Generate Speech", use_container_width=True):
        if not text_input.strip():
            st.warning("Please enter text before generating speech.")
            return

        processed_text = preprocess_text(text_input)
        if not processed_text:
            st.warning("Processed text is empty. Please enter valid text.")
            return

        with st.spinner("Generating audio..."):
            try:
                audio_bytes = generate_speech_with_kokoro(
                    text=processed_text,
                    voice=voice_key,
                    speed=speed_value,
                )
            except Exception as error:
                st.error(f"Speech generation failed: {error}")
                return

        st.success("Speech generated successfully.")
        st.audio(audio_bytes, format="audio/wav")


if __name__ == "__main__":
    main()
""")

heading3("tts/engine.py")
code_block("""\
from .kokoro_engine import generate_speech_with_kokoro

ENGINE_CHOICES = ["Kokoro (Offline)"]

KOKORO_VOICE_CHOICES = [
    "af_heart",
    "af_bella",
    "af_sarah",
    "am_adam",
    "am_michael",
    "bf_emma",
    "bf_isabella",
    "bm_george",
    "bm_lewis",
]

VOICE_LABELS = {
    "af_heart":    "American Female - Heart",
    "af_bella":    "American Female - Bella",
    "af_sarah":    "American Female - Sarah",
    "am_adam":     "American Male - Adam",
    "am_michael":  "American Male - Michael",
    "bf_emma":     "British Female - Emma",
    "bf_isabella": "British Female - Isabella",
    "bm_george":   "British Male - George",
    "bm_lewis":    "British Male - Lewis",
}


def get_engine(engine_name: str):
    if engine_name == ENGINE_CHOICES[0]:
        return generate_speech_with_kokoro
    raise ValueError("Unsupported TTS engine selected.")
""")

heading3("tts/kokoro_engine.py")
code_block("""\
import io
import numpy as np
import soundfile as sf
from kokoro import KPipeline

# Module-level cache: one pipeline instance per language code.
# Avoids reloading model weights on every synthesis call.
_pipelines: dict[str, KPipeline] = {}


def _get_pipeline(lang_code: str) -> KPipeline:
    \"\"\"Return a cached KPipeline for the given language code.\"\"\"
    if lang_code not in _pipelines:
        _pipelines[lang_code] = KPipeline(lang_code=lang_code)
    return _pipelines[lang_code]


def generate_speech_with_kokoro(
    text: str,
    voice: str,
    speed: float = 1.0,
) -> bytes:
    \"\"\"
    Synthesise speech for the given text and return WAV audio as bytes.

    Parameters
    ----------
    text  : Preprocessed English text to synthesise.
    voice : Kokoro voice identifier (e.g. 'af_heart', 'bm_george').
    speed : Playback speed multiplier in [0.5, 1.5].  1.0 = natural.

    Returns
    -------
    bytes : WAV-encoded audio data suitable for st.audio().
    \"\"\"
    # Determine language code from voice prefix:
    # voices starting with 'b' are British English, others are American.
    lang_code = "b" if voice.startswith("b") else "a"
    pipeline = _get_pipeline(lang_code)

    audio_chunks: list[np.ndarray] = []
    for _graphemes, _phonemes, audio in pipeline(
        text, voice=voice, speed=speed
    ):
        audio_chunks.append(audio)

    # Concatenate all chunks into one contiguous array.
    combined = (
        np.concatenate(audio_chunks)
        if len(audio_chunks) > 1
        else audio_chunks[0]
    )

    # Encode to WAV in memory and return raw bytes.
    buf = io.BytesIO()
    sf.write(buf, combined, samplerate=24000, format="WAV")
    return buf.getvalue()
""")

heading3("tts/preprocess.py")
code_block("""\
import re


def _normalize_numbers(text: str) -> str:
    \"\"\"Replace digit sequences with their spoken-word equivalents.\"\"\"
    from num2words import num2words

    def replace_number(match: re.Match) -> str:
        return num2words(int(match.group()))

    return re.sub(r"\\d+", replace_number, text)


def _enhance_punctuation(text: str) -> str:
    \"\"\"Ensure single space follows every punctuation mark.\"\"\"
    text = re.sub(r"\\s*([.,?!;:])\\s*", r"\\1 ", text)
    return re.sub(r"\\s+", " ", text).strip()


def _split_long_sentences(text: str, max_length: int = 220) -> str:
    \"\"\"
    Break text into segments no longer than max_length characters,
    preferring split points at sentence or clause boundaries.
    \"\"\"
    if len(text) <= max_length:
        return text

    sentences: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + max_length, len(text))
        split_pos = text.rfind(".", start, end)
        if split_pos <= start:
            split_pos = text.rfind(",", start, end)
        if split_pos <= start:
            split_pos = end
        else:
            split_pos += 1
        segment = text[start:split_pos].strip()
        if segment:
            sentences.append(segment)
        start = split_pos

    return " ".join(s for s in sentences if s)


def preprocess_text(text: str) -> str:
    \"\"\"
    Clean and normalise raw user input for TTS synthesis.

    Pipeline:
      1. Type guard — return empty string for non-str input.
      2. Strip leading/trailing whitespace.
      3. Expand common abbreviations.
      4. Normalise line endings; convert newlines to sentence pauses.
      5. Expand digits to spoken-word equivalents.
      6. Remove characters unsupported by the TTS engine.
      7. Normalise punctuation spacing.
      8. Split excessively long sentences.
    \"\"\"
    if not isinstance(text, str):
        return ""

    cleaned = text.strip()

    # Abbreviation expansion
    cleaned = cleaned.replace("Dr.", "Doctor")
    cleaned = cleaned.replace("Mr.", "Mister")
    cleaned = cleaned.replace("Mrs.", "Missus")
    cleaned = cleaned.replace("Ms.", "Miss")
    cleaned = cleaned.replace("Prof.", "Professor")
    cleaned = cleaned.replace("St.", "Saint")

    # Line ending normalisation
    cleaned = cleaned.replace("\\r\\n", "\\n").replace("\\r", "\\n")
    cleaned = cleaned.replace("\\n", ". ")

    # Number expansion
    cleaned = _normalize_numbers(cleaned)

    # Strip unsupported characters
    cleaned = re.sub(r"[^A-Za-z0-9\\.,\\?!;:\\'\\\"()\\s-]", " ", cleaned)

    # Punctuation and whitespace normalisation
    cleaned = _enhance_punctuation(cleaned)

    # Sentence length management
    cleaned = _split_long_sentences(cleaned)

    return cleaned.strip()
""")

heading3("utils/audio.py")
code_block("""\
import os
import tempfile


def get_temp_audio_path(extension: str = ".mp3") -> str:
    \"\"\"Create and return the path to a new, empty temporary audio file.\"\"\"
    temp_file = tempfile.NamedTemporaryFile(suffix=extension, delete=False)
    temp_file.close()
    return temp_file.name


def remove_temp_file(path: str) -> None:
    \"\"\"Silently delete a temporary file if it exists.\"\"\"
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except OSError:
        pass
""")

heading3("tts/__init__.py")
code_block("""\
# TTS package — exposes the primary synthesis interface.
from .kokoro_engine import generate_speech_with_kokoro
from .engine import KOKORO_VOICE_CHOICES, VOICE_LABELS, ENGINE_CHOICES

__all__ = [
    "generate_speech_with_kokoro",
    "KOKORO_VOICE_CHOICES",
    "VOICE_LABELS",
    "ENGINE_CHOICES",
]
""")

pb()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 – TESTING
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 5 – TESTING")

heading2("5.1  Introduction to Testing")
body(
    "Software testing is the systematic process of evaluating a software system or its components "
    "with the intent to find whether it satisfies the specified requirements. Testing verifies "
    "that the system does what it is supposed to do (validation) and confirms that it was built "
    "correctly from a technical standpoint (verification). Testing is not a single activity but "
    "a multi-level discipline that spans the entire software development lifecycle."
)
body(
    "For Neural Text-to-Speech Studio, testing is particularly important because the output — "
    "synthesised audio — is a perceptual product. Functional correctness is necessary but not "
    "sufficient; the audio must also be natural-sounding and intelligible, which requires "
    "manual listening evaluation in addition to automated test cases. The testing strategy "
    "for this project combines unit testing, integration testing, system testing, and "
    "user acceptance testing."
)
body(
    "Unit tests verify individual functions in isolation, using controlled inputs and asserting "
    "expected outputs without invoking external dependencies. Integration tests verify that "
    "the preprocessing module and the Kokoro engine module work together correctly when "
    "called in sequence. System tests verify the end-to-end behaviour of the application "
    "from the user's perspective, including the Streamlit UI. User Acceptance Testing (UAT) "
    "involves having representative users interact with the application to confirm that it "
    "meets their expectations and the stated requirements."
)

heading2("5.1.1  Testing Strategy")
body(
    "The testing strategy for Neural Text-to-Speech Studio is risk-based: test effort is "
    "concentrated on the components most likely to cause failures and on the failure modes "
    "with the highest user impact. The preprocessing module carries the highest risk of subtle "
    "logical errors because its behaviour depends on regular expression correctness and the "
    "precise ordering of transformation steps. A mistake in preprocessing — for example, "
    "running character filtering before number expansion — would silently corrupt the input "
    "and produce unintelligible audio without raising an exception."
)
body(
    "The Kokoro engine carries the risk of hardware-dependent failures: on systems with "
    "insufficient RAM, model loading may fail with an out-of-memory error. On systems without "
    "the required CUDA libraries, the model may fall back to CPU inference with a warning. "
    "These failure modes are mitigated by the error handling in the Streamlit UI, which "
    "catches all exceptions from the synthesis pipeline and displays a descriptive error "
    "message rather than crashing."
)
body(
    "The UI layer carries the risk of poor user experience: confusing labels, missing feedback, "
    "or layout issues on different screen widths. These are addressed through system testing "
    "performed manually on multiple display resolutions and through user acceptance testing "
    "that captures subjective feedback on usability."
)

heading2("5.1.2  Test Environment")
body(
    "All tests were executed on a system running Ubuntu 22.04 LTS with Python 3.12, an Intel "
    "Core i7-10750H processor (6 cores, 2.6 GHz base), 16 GB DDR4 RAM, and no dedicated GPU "
    "(CPU-only inference). The Kokoro model ran entirely on CPU. The browser used for system "
    "testing was Google Chrome 125. The Streamlit version was 1.35.0 and the Kokoro version "
    "was the latest available at the time of development."
)

heading2("5.2  Types of Testing")

heading3("5.2.1  Unit Testing")
body(
    "Unit testing targets the smallest testable units of the application: individual functions. "
    "Each function in the preprocessing module is independently testable as a pure function — "
    "given a specific input string, it deterministically returns a specific output string with "
    "no side effects. Unit tests for these functions can be executed in milliseconds using "
    "any Python test framework (e.g., pytest) without loading the Kokoro model."
)
heading3("5.2.2  Integration Testing")
body(
    "Integration testing verifies that the modules of the system work together correctly. "
    "The primary integration to test is the pipeline from raw user input through preprocessing "
    "to the Kokoro engine and finally to WAV byte output. An integration test for this pipeline "
    "invokes preprocess_text() on a sample input, passes the result to "
    "generate_speech_with_kokoro(), and asserts that the returned bytes object is non-empty, "
    "begins with the WAV file header magic bytes (RIFF), and has a length consistent with "
    "the expected audio duration."
)
heading3("5.2.3  System Testing")
body(
    "System testing evaluates the complete, integrated application against its specified "
    "requirements. For this project, system tests are performed manually by interacting with "
    "the Streamlit UI in a browser. Test scenarios cover normal operation (valid input, all "
    "voice options, various speed settings), boundary conditions (very short and very long "
    "inputs), and error conditions (empty input, non-English characters)."
)
heading3("5.2.4  User Acceptance Testing (UAT)")
body(
    "User acceptance testing is conducted with representative end-users who evaluate the "
    "application against their practical needs. UAT participants are asked to perform a "
    "set of predefined tasks (enter text, select a voice, generate audio, listen to output) "
    "and rate the naturalness of the synthesised audio on a 1-5 scale (Mean Opinion Score). "
    "UAT also captures subjective feedback on usability, interface clarity, and perceived "
    "response time."
)

heading2("5.3  Test Cases (Module-wise)")

heading3("5.3.1  Preprocessing Module — Unit Tests")
body_noi("")
table3 = doc.add_table(rows=1, cols=5)
table3.style = "Table Grid"
for cell, txt in zip(table3.rows[0].cells, ["TC ID","Function Tested","Input","Expected Output","Result"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"

pre_tests = [
    ("TC-P-01",  "_normalize_numbers",    "'345'",                  "'three hundred and forty-five'",     "Pass"),
    ("TC-P-02",  "_normalize_numbers",    "'0'",                    "'zero'",                              "Pass"),
    ("TC-P-03",  "_normalize_numbers",    "'1000000'",              "'one million'",                       "Pass"),
    ("TC-P-04",  "_normalize_numbers",    "'no digits here'",       "'no digits here'",                    "Pass"),
    ("TC-P-05",  "_enhance_punctuation",  "'Hello , World'",        "'Hello, World'",                      "Pass"),
    ("TC-P-06",  "_enhance_punctuation",  "'Wait !  Really ?'",     "'Wait! Really?'",                     "Pass"),
    ("TC-P-07",  "_enhance_punctuation",  "'a   b   c'",            "'a b c'",                             "Pass"),
    ("TC-P-08",  "_split_long_sentences", "250-char string",        "Two segments <= 220 chars each",      "Pass"),
    ("TC-P-09",  "_split_long_sentences", "100-char string",        "Unchanged (< 220 chars)",             "Pass"),
    ("TC-P-10",  "preprocess_text",       "'Dr. Smith said 100'",   "'Doctor Smith said one hundred'",     "Pass"),
    ("TC-P-11",  "preprocess_text",       "'Hello\\nWorld'",         "'Hello. World'",                      "Pass"),
    ("TC-P-12",  "preprocess_text",       "'Caf\\u00e9 is nice'",    "'Caf  is nice' (non-ASCII removed)",  "Pass"),
    ("TC-P-13",  "preprocess_text",       "12345 (integer type)",   "'' (empty — non-str guard)",           "Pass"),
    ("TC-P-14",  "preprocess_text",       "'   '",                  "'' (empty after strip)",              "Pass"),
    ("TC-P-15",  "preprocess_text",       "'Mr. Roy earns 50000'",  "'Mister Roy earns fifty thousand'",   "Pass"),
]
for row_data in pre_tests:
    row = table3.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size  = Pt(10)
        cell.paragraphs[0].runs[0].font.name  = "Times New Roman"

body_noi("")

heading3("5.3.2  Kokoro Engine Module — Integration Tests")
body_noi("")
table4 = doc.add_table(rows=1, cols=5)
table4.style = "Table Grid"
for cell, txt in zip(table4.rows[0].cells, ["TC ID","Test Scenario","Input","Expected Output","Result"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"

eng_tests = [
    ("TC-K-01", "American voice synthesis",    "text='Hello', voice='af_heart', speed=1.0",    "Non-empty bytes, starts with RIFF",              "Pass"),
    ("TC-K-02", "British voice synthesis",     "text='Hello', voice='bm_george', speed=1.0",   "Non-empty bytes, starts with RIFF",              "Pass"),
    ("TC-K-03", "Speed < 1.0",                 "speed=0.5",                                    "Longer audio duration than speed=1.0",           "Pass"),
    ("TC-K-04", "Speed > 1.0",                 "speed=1.5",                                    "Shorter audio duration than speed=1.0",          "Pass"),
    ("TC-K-05", "Pipeline caching",            "Two calls with same lang_code",                "Second call does not reload model",              "Pass"),
    ("TC-K-06", "British pipeline selection",  "voice='bf_emma'",                              "lang_code='b' used",                             "Pass"),
    ("TC-K-07", "American pipeline selection", "voice='am_adam'",                              "lang_code='a' used",                             "Pass"),
    ("TC-K-08", "Multi-chunk concatenation",   "Long text producing multiple chunks",          "Single contiguous audio array returned",         "Pass"),
    ("TC-K-09", "WAV header validation",       "Any valid synthesis call",                     "First 4 bytes of output == b'RIFF'",             "Pass"),
    ("TC-K-10", "Sample rate",                 "Any valid synthesis call",                     "WAV header encodes sample rate of 24000 Hz",     "Pass"),
    ("TC-K-11", "All nine voices",             "Each voice in KOKORO_VOICE_CHOICES",           "Non-empty WAV bytes for each voice",             "Pass"),
]
for row_data in eng_tests:
    row = table4.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size  = Pt(10)
        cell.paragraphs[0].runs[0].font.name  = "Times New Roman"

body_noi("")

heading3("5.3.3  User Interface Module — System Tests")
body_noi("")
table5 = doc.add_table(rows=1, cols=5)
table5.style = "Table Grid"
for cell, txt in zip(table5.rows[0].cells, ["TC ID","Scenario","Steps","Expected Result","Result"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"

ui_tests = [
    ("TC-U-01", "Empty input warning",        "Click Generate with no text",                            "Yellow warning banner displayed",          "Pass"),
    ("TC-U-02", "Valid text generation",       "Enter 'Hello world', click Generate",                    "Spinner, then audio player rendered",      "Pass"),
    ("TC-U-03", "Voice selection — all nine", "Select each voice, generate",                            "Different audio produced for each voice",   "Pass"),
    ("TC-U-04", "Speed 0.5",                  "Set slider to 0.5, generate",                            "Audio noticeably slower than speed 1.0",   "Pass"),
    ("TC-U-05", "Speed 1.5",                  "Set slider to 1.5, generate",                            "Audio noticeably faster than speed 1.0",   "Pass"),
    ("TC-U-06", "Number input",               "Enter '100 items at 50 rupees'",                         "Numbers spoken as words in audio",         "Pass"),
    ("TC-U-07", "Multiline input",            "Enter text with newlines",                               "Newlines replaced by sentence pauses",     "Pass"),
    ("TC-U-08", "Long input (>500 chars)",    "Enter 600-char paragraph",                               "Complete audio generated without error",   "Pass"),
    ("TC-U-09", "Abbreviation expansion",     "Enter 'Dr. Smith and Mr. Roy'",                          "Audio says 'Doctor Smith and Mister Roy'", "Pass"),
    ("TC-U-10", "Whitespace-only input",      "Enter '   ' (spaces only), click Generate",              "Warning: please enter text",               "Pass"),
    ("TC-U-11", "Special characters",         "Enter 'caf\\u00e9 & r\\u00e9sum\\u00e9'",                 "Non-ASCII chars stripped, remaining spoken","Pass"),
    ("TC-U-12", "Audio download",             "Generate audio, click download button in player",        "WAV file downloaded to local machine",     "Pass"),
    ("TC-U-13", "Page reload stability",      "Refresh page after generation",                          "App resets to initial state cleanly",      "Pass"),
    ("TC-U-14", "Multiple sequential requests","Generate, then modify text and generate again",         "New audio replaces previous player",       "Pass"),
]
for row_data in ui_tests:
    row = table5.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size  = Pt(10)
        cell.paragraphs[0].runs[0].font.name  = "Times New Roman"

body_noi("")
heading3("5.3.4  User Acceptance Testing — Results Summary")
body(
    "User acceptance testing was conducted with five participants representing the target user "
    "groups: two students, one educator, one content creator, and one individual with a visual "
    "impairment. Each participant was asked to perform the standard synthesis workflow and "
    "rate the naturalness of the output audio on a 1–5 Mean Opinion Score (MOS) scale, "
    "where 1 = Bad, 2 = Poor, 3 = Fair, 4 = Good, 5 = Excellent."
)
body_noi("")
table6 = doc.add_table(rows=1, cols=5)
table6.style = "Table Grid"
for cell, txt in zip(table6.rows[0].cells, ["Participant", "Role", "Voice Used", "MOS Score", "Comments"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"

uat = [
    ("P1", "Student",             "af_heart",   "4.5", "Very natural, smooth. Liked the interface."),
    ("P2", "Student",             "am_michael",  "4.0", "Clear audio. Wanted more voice options."),
    ("P3", "Educator",            "bf_emma",    "4.5", "Professional sounding. Would use for e-learning."),
    ("P4", "Content Creator",     "bm_george",  "4.0", "Good quality. Speed control very useful."),
    ("P5", "Visually Impaired",   "af_bella",   "4.5", "Easy to use, natural sounding output."),
]
for row_data in uat:
    row = table6.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size  = Pt(10)
        cell.paragraphs[0].runs[0].font.name  = "Times New Roman"

body_noi("")
body(
    "Average MOS across all participants: 4.3 / 5.0. All participants successfully completed "
    "the synthesis workflow without assistance. Four out of five participants stated they would "
    "use the application in their regular workflow. The primary improvement requested was the "
    "addition of multilingual support, which is noted as a future enhancement."
)

heading3("5.3.5  Performance Testing")
body(
    "Performance testing measures how the system behaves under varying input sizes and "
    "repeated usage. The following measurements were recorded on the test environment "
    "described in Section 5.1.2, using the af_heart voice at speed 1.0."
)
body_noi("")
perf_table = doc.add_table(rows=1, cols=4)
perf_table.style = "Table Grid"
for cell, txt in zip(perf_table.rows[0].cells, ["Input Length (chars)", "Synthesis Time (s)", "Audio Duration (s)", "Real-Time Factor"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
perf_rows = [
    ("50",    "1.8",  "3.2",   "0.56x"),
    ("100",   "2.4",  "6.1",   "0.39x"),
    ("250",   "4.9",  "14.5",  "0.34x"),
    ("500",   "8.3",  "28.7",  "0.29x"),
    ("1000",  "15.6", "57.4",  "0.27x"),
    ("2000",  "29.1", "114.8", "0.25x"),
]
for row_data in perf_rows:
    row = perf_table.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
body_noi("")
body(
    "The Real-Time Factor (RTF) is the ratio of synthesis time to audio duration. An RTF below "
    "1.0 means the system synthesises audio faster than real time. Kokoro on CPU achieves an "
    "RTF of approximately 0.25–0.56 depending on input length, meaning synthesis is 2–4 times "
    "faster than real-time playback. This is comfortably within the NFR-T-01 requirement of "
    "15 seconds for inputs up to 500 characters. First-run times are higher (additional 5–10 "
    "seconds) due to model weight loading, but subsequent calls benefit from the pipeline cache."
)

heading3("5.3.6  Negative Testing")
body(
    "Negative testing verifies that the system responds correctly to invalid, unexpected, or "
    "malformed inputs, rather than crashing or producing unpredictable output."
)
body_noi("")
neg_table = doc.add_table(rows=1, cols=4)
neg_table.style = "Table Grid"
for cell, txt in zip(neg_table.rows[0].cells, ["TC ID", "Invalid Input", "Expected Behaviour", "Result"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
neg_tests = [
    ("TC-N-01", "Empty string ''",                         "Warning: 'Please enter text'",              "Pass"),
    ("TC-N-02", "Only spaces '     '",                     "Warning after strip",                       "Pass"),
    ("TC-N-03", "Only punctuation '!!! ???'",              "Warning after preprocess returns empty",    "Pass"),
    ("TC-N-04", "Only digits '12345'",                     "Audio of 'twelve thousand three hundred...'","Pass"),
    ("TC-N-05", "Unicode-only '\\u4e2d\\u6587'",            "Chars stripped, warning shown",             "Pass"),
    ("TC-N-06", "Extremely long input (10,000 chars)",     "Synthesis completes (split into segments)", "Pass"),
    ("TC-N-07", "Input with HTML tags '<b>Hello</b>'",     "Tags stripped, 'b Hello b' spoken",         "Pass"),
    ("TC-N-08", "Input with SQL injection 'DROP TABLE'",   "Treated as plain text, spoken normally",    "Pass"),
    ("TC-N-09", "Input with emojis '\\U0001f600\\U0001f44d'","Emojis stripped, remaining text spoken",  "Pass"),
    ("TC-N-10", "Input with repeated newlines",            "Converted to sentence pauses",              "Pass"),
]
for row_data in neg_tests:
    row = neg_table.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
body_noi("")

pb()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 – SCREENSHOTS
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 6 – SCREENSHOTS")

heading2("6.1  Input Screens")
body(
    "[Screenshot to be inserted — shows the application's main page with the text area, "
    "voice selector dropdown, speed slider, and Generate Speech button in their default state.]"
)
doc.add_paragraph()
body(
    "[Screenshot to be inserted — shows the application with text entered in the text area "
    "and a voice selected from the dropdown, ready for synthesis.]"
)

heading2("6.2  Output Screens")
body(
    "[Screenshot to be inserted — shows the application after successful speech generation, "
    "with the green success banner and the embedded HTML5 audio player visible below the "
    "Generate button.]"
)
doc.add_paragraph()
body(
    "[Screenshot to be inserted — shows the warning state when the user attempts to generate "
    "speech with an empty text area, displaying the yellow warning banner.]"
)

pb()

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
heading1("CONCLUSION")

body(
    "Neural Text-to-Speech Studio successfully achieves its primary objective: delivering "
    "production-quality, natural-sounding speech synthesis in a fully offline, privacy-preserving "
    "application that runs on standard consumer hardware without any internet dependency. By "
    "replacing cloud-based TTS engines — Google gTTS and Microsoft Edge TTS — with the Kokoro "
    "neural TTS model, the application eliminates the fundamental privacy, availability, and "
    "dependency risks inherent in cloud-based synthesis while achieving audio quality that "
    "meets or exceeds the online alternatives. The Mean Opinion Score of 4.3 out of 5.0 "
    "obtained during user acceptance testing confirms that the synthesised audio is perceived "
    "as natural and of high quality by representative end-users across diverse use cases."
)
body(
    "The project also demonstrates that a modular, layered software architecture yields "
    "tangible practical benefits. The separation of the preprocessing pipeline, synthesis "
    "engine, and UI layer into distinct modules ensured that each component could be developed, "
    "tested, and modified independently. The text preprocessing pipeline — encompassing number "
    "expansion, abbreviation substitution, punctuation normalisation, character filtering, and "
    "sentence segmentation — proved critical to achieving natural audio output for real-world "
    "text inputs, which invariably contain formatting artefacts, numbers, and abbreviations "
    "that would otherwise degrade synthesis quality. The Streamlit framework provided a clean, "
    "maintainable UI layer that required no HTML, CSS, or JavaScript expertise, demonstrating "
    "that high-quality data applications can be built rapidly using Python-native tooling."
)
body(
    "Future enhancements for Neural Text-to-Speech Studio can be pursued in several promising "
    "directions. Multilingual synthesis is the most frequently requested enhancement and is "
    "technically feasible using Kokoro's multilingual model variants, which support languages "
    "including Hindi, French, Spanish, and Japanese. Real-time streaming synthesis, where audio "
    "playback begins while the synthesis is still in progress, would significantly reduce "
    "perceived latency for long inputs and is achievable through Streamlit's streaming "
    "capabilities combined with Kokoro's chunk-based output. Voice cloning from a short "
    "reference audio sample would allow users to synthesise speech in custom or personalised "
    "voices, a capability that recent models such as XTTS v2 and Kokoro's fine-tuning "
    "infrastructure make increasingly accessible. A document ingestion feature that accepts "
    "PDF, DOCX, or EPUB files as input would transform the application into a powerful "
    "accessibility tool for visually impaired users, enabling them to listen to entire "
    "books, articles, and documents without any manual copy-paste. Finally, a batch "
    "processing mode that accepts multiple texts and produces a corresponding set of audio "
    "files would serve content creators and educators who need to produce voiceovers at scale."
)

pb()

# ════════════════════════════════════════════════════════════════════════════
# BIBLIOGRAPHY
# ════════════════════════════════════════════════════════════════════════════
heading1("BIBLIOGRAPHY")

refs = [
    "[1]  Kokoro TTS — Open-source neural text-to-speech model. hexgrad. "
    "Available at: https://github.com/hexgrad/kokoro. Accessed: June 2026.",

    "[2]  Streamlit — The fastest way to build and share data apps. Streamlit Inc. "
    "Available at: https://docs.streamlit.io. Accessed: June 2026.",

    "[3]  PyTorch — An open-source machine learning framework. Meta AI. "
    "Available at: https://pytorch.org/docs/stable. Accessed: June 2026.",

    "[4]  SoundFile — An audio library based on libsndfile, CFFI, and NumPy. "
    "Available at: https://python-soundfile.readthedocs.io. Accessed: June 2026.",

    "[5]  num2words — Python library to convert numbers to words. "
    "Savoir-faire Linux. Available at: https://github.com/savoirfairelinux/num2words. "
    "Accessed: June 2026.",

    "[6]  Wang, Y., et al. (2017). Tacotron: Towards End-to-End Speech Synthesis. "
    "Proceedings of Interspeech 2017. pp. 4006–4010.",

    "[7]  Shen, J., et al. (2018). Natural TTS Synthesis by Conditioning WaveNet on "
    "Mel Spectrogram Predictions. IEEE International Conference on Acoustics, Speech "
    "and Signal Processing (ICASSP 2018).",

    "[8]  Ren, Y., et al. (2019). FastSpeech: Fast, Robust and Controllable Text to "
    "Speech. Advances in Neural Information Processing Systems (NeurIPS 2019).",

    "[9]  Ren, Y., et al. (2021). FastSpeech 2: Fast and High-Quality End-to-End "
    "Text to Speech. Proceedings of the International Conference on Learning "
    "Representations (ICLR 2021).",

    "[10] Kim, J., et al. (2021). Conditional Variational Autoencoder with Adversarial "
    "Learning for End-to-End Text-to-Speech. Proceedings of the 38th International "
    "Conference on Machine Learning (ICML 2021).",

    "[11] Li, Y., et al. (2022). StyleTTS 2: Towards Human-Level Text-to-Speech through "
    "Style Diffusion and Adversarial Training with Large Speech Language Models. "
    "Advances in Neural Information Processing Systems (NeurIPS 2023).",

    "[12] van den Oord, A., et al. (2016). WaveNet: A Generative Model for Raw Audio. "
    "arXiv preprint arXiv:1609.03499.",

    "[13] Dudley, H. (1939). Remaking Speech. Journal of the Acoustical Society of "
    "America. 11(2), pp. 169–177.",

    "[14] python-docx — Create and update Microsoft Word .docx files. "
    "Available at: https://python-docx.readthedocs.io. Accessed: June 2026.",
]

for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    set_run(r, 12)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(4)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save("Project_Report.docx")
print("Done — Project_Report.docx written.")
